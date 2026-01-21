from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class QuestionPaperGenerator:
    """Generate Question Paper in Word format."""
    
    def __init__(self, college_name, exam_name, date, questions):
        """
        Initialize the MCQ Paper Generator.
        
        Args:
            college_name: Name of the college
            exam_name: Name of the exam
            date: Date of the exam
            questions: Dictionary of questions {id: question_data}
        """
        self.college_name = college_name
        self.exam_name = exam_name
        self.date = date
        self.questions = questions
        self.doc = None
    
    def add_header(self):
        """Add header with college name and exam details."""
        # College name as main heading
        heading = self.doc.add_heading(self.college_name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_format = heading.runs[0].font
        heading_format.size = Pt(18)
        heading_format.bold = True
        
        # Add a line
        self.doc.add_paragraph()
        
        # Exam details - side by side format
        details_table = self.doc.add_table(rows=2, cols=2)
        details_table.autofit = False
        
        # Date
        date_cell = details_table.cell(0, 0)
        date_para = date_cell.paragraphs[0]
        date_para.text = f"Date: {self.date}"
        date_run = date_para.runs[0]
        date_run.font.size = Pt(12)
        date_run.font.bold = True
        
        # Exam Name
        exam_cell = details_table.cell(0, 1)
        exam_para = exam_cell.paragraphs[0]
        exam_para.text = f"Exam: {self.exam_name}"
        exam_run = exam_para.runs[0]
        exam_run.font.size = Pt(12)
        exam_run.font.bold = True
        
        # Remove borders
        for row in details_table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is None:
                    from docx.oxml import parse_xml
                    tcBorders = parse_xml(
                        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        '<w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/>'
                        '<w:insideH w:val="none"/><w:insideV w:val="none"/></w:tcBorders>'
                    )
                    tcPr.append(tcBorders)
        
        self.doc.add_paragraph()
    
    def add_instructions(self):
        """Add instructions section."""
        instructions_heading = self.doc.add_heading('Instructions:', level=2)
        instructions_heading.runs[0].font.size = Pt(12)
        
        instructions = [
            "Attempt all questions.",
            "Each question carries one mark.",
            "There is no negative marking.",
            "Select the most appropriate option from A, B, C, and D."
        ]
        
        for instruction in instructions:
            self.doc.add_paragraph(instruction, style='List Bullet')
        
        self.doc.add_paragraph()
    
    def add_questions(self):
        """Add questions to the document."""
        for question_id in sorted(self.questions.keys()):
            question_data = self.questions[question_id]
            
            # Question number and text
            question_para = self.doc.add_paragraph()
            question_para.paragraph_format.left_indent = Inches(0)
            question_run = question_para.add_run(f"Q{question_id}. {question_data['question']}")
            question_run.font.bold = True
            question_run.font.size = Pt(11)
            
            # Options
            for option_key in ['A', 'B', 'C', 'D']:
                option_text = question_data['options'][option_key]
                option_para = self.doc.add_paragraph()
                option_para.paragraph_format.left_indent = Inches(0.25)
                option_run = option_para.add_run(f"{option_key}) {option_text}")
                option_run.font.size = Pt(10)
            
            # Add space between questions
            self.doc.add_paragraph()
    
    def generate(self, output_path):
        """Generate the question paper document.
        
        Args:
            output_path: Path to save the Word document
        """
        self.doc = Document()
        
        self.add_header()
        self.add_instructions()
        self.add_questions()
        
        self.doc.save(output_path)
        return output_path


class AnswerKeyGenerator:
    """Generate Answer Key and Solutions in Word format."""
    
    def __init__(self, college_name, exam_name, date, questions):
        """
        Initialize the Answer Key Generator.
        
        Args:
            college_name: Name of the college
            exam_name: Name of the exam
            date: Date of the exam
            questions: Dictionary of questions {id: question_data}
        """
        self.college_name = college_name
        self.exam_name = exam_name
        self.date = date
        self.questions = questions
        self.doc = None
    
    def add_header(self):
        """Add header with college name and exam details."""
        # College name as main heading
        heading = self.doc.add_heading(f"{self.college_name} - Answer Key & Solutions", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_format = heading.runs[0].font
        heading_format.size = Pt(16)
        heading_format.bold = True
        
        # Add a line
        self.doc.add_paragraph()
        
        # Exam details - side by side format
        details_table = self.doc.add_table(rows=1, cols=2)
        
        # Date
        date_cell = details_table.cell(0, 0)
        date_para = date_cell.paragraphs[0]
        date_para.text = f"Date: {self.date}"
        date_run = date_para.runs[0]
        date_run.font.size = Pt(11)
        date_run.font.bold = True
        
        # Exam Name
        exam_cell = details_table.cell(0, 1)
        exam_para = exam_cell.paragraphs[0]
        exam_para.text = f"Exam: {self.exam_name}"
        exam_run = exam_para.runs[0]
        exam_run.font.size = Pt(11)
        exam_run.font.bold = True
        
        # Remove borders
        for row in details_table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is None:
                    from docx.oxml import parse_xml
                    tcBorders = parse_xml(
                        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        '<w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/>'
                        '<w:insideH w:val="none"/><w:insideV w:val="none"/></w:tcBorders>'
                    )
                    tcPr.append(tcBorders)
        
        self.doc.add_paragraph()
    
    def add_answer_key_table(self):
        """Add answer key table."""
        ans_heading = self.doc.add_heading('Answer Key:', level=2)
        ans_heading.runs[0].font.size = Pt(12)
        
        # Create table for answers
        num_questions = len(self.questions)
        num_cols = 10  # 10 columns for better layout
        num_rows = (num_questions + num_cols - 1) // num_cols + 1
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        for i in range(num_cols):
            table.cell(0, i).text = ""
        
        # Fill in answers
        question_ids = sorted(self.questions.keys())
        for idx, qid in enumerate(question_ids):
            row = (idx // num_cols) + 1
            col = idx % num_cols
            correct_ans = self.questions[qid]['correct_answer']
            
            cell = table.cell(row, col)
            cell.text = f"Q{qid}: {correct_ans}"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.bold = True
        
        self.doc.add_paragraph()
    
    def add_detailed_solutions(self):
        """Add detailed solutions for each question."""
        sol_heading = self.doc.add_heading('Detailed Solutions:', level=2)
        sol_heading.runs[0].font.size = Pt(12)
        
        for question_id in sorted(self.questions.keys()):
            question_data = self.questions[question_id]
            correct_ans = question_data['correct_answer']
            explanation = question_data['explanation']
            
            # Question
            q_para = self.doc.add_paragraph()
            q_run = q_para.add_run(f"Q{question_id}. {question_data['question']}")
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            
            # Answer
            ans_para = self.doc.add_paragraph()
            ans_para.paragraph_format.left_indent = Inches(0.25)
            ans_run = ans_para.add_run(f"Answer: {correct_ans}")
            ans_run.font.bold = True
            ans_run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
            ans_run.font.size = Pt(10)
            
            # Explanation
            exp_para = self.doc.add_paragraph()
            exp_para.paragraph_format.left_indent = Inches(0.25)
            exp_run = exp_para.add_run(f"Explanation: {explanation}")
            exp_run.font.size = Pt(10)
            exp_run.font.italic = True
            
            # Add space between questions
            self.doc.add_paragraph()
    
    def generate(self, output_path):
        """Generate the answer key document.
        
        Args:
            output_path: Path to save the Word document
        """
        self.doc = Document()
        
        self.add_header()
        self.add_answer_key_table()
        self.add_detailed_solutions()
        
        self.doc.save(output_path)
        return output_path
